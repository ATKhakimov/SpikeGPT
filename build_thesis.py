"""Build thesis DOCX following the MIET/NIIME thesis format."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

RU_TITLE = (
    "Влияние морфологической сложности естественного языка "
    "на нейроморфную спарсити импульсных языковых моделей"
)
EN_TITLE = (
    "Impact of natural language morphological complexity "
    "on neuromorphic sparsity of spiking language models"
)

RU_AUTHORS = "Хакимов Н. Д.\u00B9"
EN_AUTHORS = "Khakimov N. D.\u00B9"

RU_AFFIL = [
    "\u00B9 Национальный исследовательский университет "
    "«Московский физико-технический институт»",
    "141701, Московская область, г. Долгопрудный, Институтский переулок, д. 9.",
]
EN_AFFIL = [
    "\u00B9 Moscow Institute of Physics and Technology",
    "141701, Moscow region, Dolgoprudny, Institutskiy pereulok, 9.",
]

EMAIL = "nodosty@phystech.edu"

RU_ABSTRACT = (
    "Предложена методология количественной оценки влияния морфологической "
    "сложности языка на спайковую активность нейроморфных языковых моделей. "
    "Русский язык требует на 53 % больше спайков, чем английский."
)

RU_KEYWORDS = (
    "импульсные нейронные сети; нейроморфные вычисления; SpikeGPT; "
    "LIF-нейроны; морфология языка."
)

RU_BODY = [
    (
        "Импульсные нейронные сети (ИНС) — энергоэффективная альтернатива "
        "классическим нейросетям: энергопотребление нейроморфных ускорителей "
        "(Intel Loihi, BrainScaleS) пропорционально средней частоте спайков "
        "[1]. Модель SpikeGPT [2] на архитектуре RWKV с бинарными "
        "LIF-нейронами впервые показала применимость ИНС к генеративному "
        "языковому моделированию, однако обучена только на английском, а её "
        "способности на русском околонулевые. Открытые русскоязычные "
        "GPT-модели семейства ruGPT-3 не получили широкого распространения "
        "и не имеют нейроморфных аналогов, что создаёт разрыв в "
        "исследовательской инфраструктуре для русского языка."
    ),
    (
        "В работе с нуля обучена русскоязычная SpikeGPT (12 слоёв, "
        "d_model = 512, ~100 млн параметров) на корпусе Тайга (~1,8 млрд "
        "токенов, BPE ruGPT-3) на NVIDIA A100 до валидационной перплексии "
        "59,79 и на её основе проведено послойное сравнение спайковой "
        "активности с открытой SpikeGPT-OpenWebText-216M [2] на ~1000 "
        "фрагментах из Тайги и OpenWebText."
    ),
    (
        "Научная новизна. Получена первая работоспособная русскоязычная "
        "импульсная языковая модель на архитектуре SpikeGPT и впервые "
        "показано, что морфологическая сложность языка — значимый фактор "
        "энергоэффективности ИНС: средний firing rate 33,2 % против 21,7 % у "
        "английской (рост 53 %). Различие локализовано в средних слоях: "
        "англоязычная модель почти неактивна (<5 %), русскоязычная "
        "поддерживает 20–25 %, что связано с ранней обработкой падежных и "
        "согласовательных зависимостей."
    ),
    (
        "Обученная модель и методология воспроизводимы и открыты для "
        "сообщества: это снимает барьер входа для дальнейших исследований "
        "спайковых языковых моделей на русском и иных морфологически "
        "богатых языках (немецкий, финский, арабский) — от доменного "
        "дообучения до запуска на нейроморфном железе — и позволяет "
        "учитывать язык целевого применения при проектировании "
        "энергобюджета нейроморфных систем."
    ),
]

EN_ABSTRACT = (
    "A methodology for quantitative evaluation of the impact of language "
    "morphological complexity on spiking activity of neuromorphic language "
    "models is proposed. Russian is shown to require 53 % more spikes than "
    "English."
)
EN_KEYWORDS = (
    "spiking neural networks; neuromorphic computing; SpikeGPT; LIF neurons; "
    "language morphology."
)

EN_BODY = [
    (
        "Spiking neural networks (SNNs) are an energy-efficient alternative "
        "to conventional networks: the energy consumption of neuromorphic "
        "accelerators (Intel Loihi, BrainScaleS) is proportional to the "
        "average spike rate [1]. The SpikeGPT model [2] on the RWKV "
        "architecture with binary LIF neurons was the first to demonstrate "
        "the applicability of SNNs to generative language modelling, yet it "
        "is trained on English only and its Russian capabilities are "
        "negligible. Open Russian GPT-like models of the ruGPT-3 family "
        "have not gained wide adoption and have no neuromorphic "
        "counterparts, creating a gap in the research infrastructure for "
        "Russian."
    ),
    (
        "In this work a Russian SpikeGPT (12 layers, d_model = 512, ~100M "
        "parameters) is trained from scratch on the Taiga corpus (~1.8B "
        "tokens, ruGPT-3 BPE) on an NVIDIA A100 down to a validation "
        "perplexity of 59.79, and on its basis a per-layer comparison of "
        "spiking activity against the open SpikeGPT-OpenWebText-216M [2] is "
        "performed on ~1000 fragments from Taiga and OpenWebText."
    ),
    (
        "Scientific novelty. The first operational Russian-language spiking "
        "language model based on the SpikeGPT architecture is obtained, and "
        "it is shown for the first time that language morphological "
        "complexity is a significant factor of SNN energy efficiency: the "
        "mean firing rate is 33.2 % versus 21.7 % for the English model "
        "(a 53 % increase). The difference is localised in the middle "
        "layers: the English model is nearly silent (<5 %) while the "
        "Russian one sustains 20–25 %, presumably due to early processing "
        "of case and agreement dependencies."
    ),
    (
        "The trained model and the methodology are reproducible and open to "
        "the community: this lowers the entry barrier for further research "
        "on spiking language models in Russian and in other morphologically "
        "rich languages (German, Finnish, Arabic) — from domain-specific "
        "fine-tuning to deployment on neuromorphic hardware — and allows "
        "the target language to be taken into account when designing the "
        "energy budget of neuromorphic systems."
    ),
]

REFERENCES = [
    "Davies M. [et al.] Loihi: A neuromorphic manycore processor with on-chip "
    "learning // IEEE Micro. 2018. V. 38. № 1. pp. 82–99.",
    "Zhu R.-J. [et al.] SpikeGPT: Generative pre-trained language model with "
    "spiking neural networks // arXiv preprint arXiv:2302.13939. 2023.",
    "Shavrina T., Shapovalova O. «Taiga» syntax tree corpus and parser // "
    "Proceedings of CORPORA 2017. 2017.",
]


def set_times(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_para(doc, text, *, bold=False, italic=False, align=None, size=12):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_times(r, size=size, bold=bold, italic=italic)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_times(r, size=12)


def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # --- Russian block ---
    add_para(doc, "УДК 004.852:004.032.26      DOI:", bold=True)
    add_para(doc, "")
    add_para(doc, RU_TITLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "")
    add_para(doc, RU_AUTHORS, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "")
    for line in RU_AFFIL:
        add_para(doc, line, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, EMAIL, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "")
    add_para(doc, "Аннотация", bold=True)
    add_body(doc, RU_ABSTRACT)
    add_para(doc, "")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run("Ключевые слова: ")
    set_times(r1, bold=True)
    r2 = p.add_run(RU_KEYWORDS)
    set_times(r2)

    for para in RU_BODY:
        add_body(doc, para)

    add_para(doc, "")

    # --- English block ---
    add_para(doc, EN_TITLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "")
    add_para(doc, EN_AUTHORS, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "")
    for line in EN_AFFIL:
        add_para(doc, line, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, EMAIL, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "")
    add_para(doc, "Abstract", bold=True)
    add_body(doc, EN_ABSTRACT)
    add_para(doc, "")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run("Keywords: ")
    set_times(r1, bold=True)
    r2 = p.add_run(EN_KEYWORDS)
    set_times(r2)

    for para in EN_BODY:
        add_body(doc, para)

    add_para(doc, "")
    add_para(doc, "Литература", bold=True, italic=True)
    for i, ref in enumerate(REFERENCES, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"{i}. {ref}")
        set_times(r, size=12)

    out = "thesis.docx"
    doc.save(out)

    # Count characters (Russian text + English text + references)
    ru_chars = len(RU_ABSTRACT) + len(RU_KEYWORDS) + sum(len(p) for p in RU_BODY)
    en_chars = len(EN_ABSTRACT) + len(EN_KEYWORDS) + sum(len(p) for p in EN_BODY)
    refs_chars = sum(len(r) for r in REFERENCES)
    print(f"RU body+abstract+keywords: {ru_chars}")
    print(f"EN body+abstract+keywords: {en_chars}")
    print(f"References:                {refs_chars}")
    print(f"TOTAL (RU+EN+refs):        {ru_chars + en_chars + refs_chars}")
    print(f"RU abstract length:        {len(RU_ABSTRACT)}  (limit: 300)")
    print(f"Saved to {out}")


if __name__ == "__main__":
    main()
