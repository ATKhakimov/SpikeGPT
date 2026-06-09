#!/usr/bin/env python3
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
DEFAULT_MD = ROOT / "ARTICLE" / "spikerugpt_conference_article_draft.md"
DEFAULT_DOCX = ROOT / "ARTICLE" / "spikerugpt_conference_article_draft.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text.strip())
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(0)

    for name, size in [("Title", 16), ("Heading 1", 14), ("Heading 2", 14)]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(6)

    for name in ["List Number", "List Bullet"]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)
        style.paragraph_format.left_indent = Cm(1.25)
        style.paragraph_format.first_line_indent = Cm(0)


def add_inline_markdown(paragraph, text: str) -> None:
    text = text.replace("**", "")
    text = text.replace("`", "")
    pos = 0
    for match in re.finditer(r"\^([^\\^]+)\^", text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        run = paragraph.add_run(match.group(1))
        run.font.superscript = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)


def add_paragraph(doc: Document, text: str, *, style: str | None = None, center: bool = False, first_line: bool = True) -> None:
    p = doc.add_paragraph(style=style)
    if not first_line:
        p.paragraph_format.first_line_indent = Cm(0)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
    add_inline_markdown(p, text)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines: list[str] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        table_lines.append(lines[i].strip())
        i += 1

    rows: list[list[str]] = []
    for idx, line in enumerate(table_lines):
        cells = [c.strip() for c in line.strip("|").split("|")]
        if idx == 1 and all(set(c.replace(":", "").strip()) <= {"-"} for c in cells):
            continue
        rows.append(cells)
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, text, bold=(r_idx == 0), size=9 if cols > 5 else 10)
            if r_idx == 0:
                set_cell_shading(cell, "D9EAF7")
    doc.add_paragraph()


def add_code_block(doc: Document, block: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.5)
    for idx, line in enumerate(block):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(10)


def add_image(doc: Document, md_path: Path, rel: str) -> None:
    image_path = (md_path.parent / rel).resolve()
    if not image_path.exists():
        add_paragraph(doc, f"[Изображение не найдено: {rel}]", first_line=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.1))


def build_docx(md_path: Path = DEFAULT_MD, output_path: Path = DEFAULT_DOCX) -> None:
    doc = Document()
    configure_document(doc)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_block: list[str] = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_block)
                code_block = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_block.append(raw)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        image_match = re.match(r"!\[[^\]]*\]\(([^)]+)\)", stripped)
        if image_match:
            add_image(doc, md_path, image_match.group(1))
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        if stripped.startswith("# "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            add_inline_markdown(p, text)
            i += 1
            continue

        if stripped.startswith("## "):
            text = stripped[3:].strip()
            add_paragraph(doc, text, style="Heading 1", first_line=False)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            add_paragraph(doc, re.sub(r"^\d+\.\s+", "", stripped), style="List Number", first_line=False)
            i += 1
            continue

        if stripped.startswith("- "):
            add_paragraph(doc, stripped[2:], style="List Bullet", first_line=False)
            i += 1
            continue

        is_caption = stripped.startswith("Рис. ")
        is_author_block = i < 8 and not stripped.startswith("УДК")
        add_paragraph(doc, stripped, center=is_caption or is_author_block, first_line=not (is_caption or is_author_block))
        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    build_docx()
